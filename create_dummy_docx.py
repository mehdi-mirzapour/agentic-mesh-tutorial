from docx import Document
import os

def create_dummy_docx(filename="dummy_test.docx"):
    doc = Document()
    doc.add_heading("Mesh Test Document", 0)
    
    paragraphs = [
        "This is the first paragraph. It contains some basic information.",
        "The second paragraph is a bit more complex. It has multiple sentences. Ideally, the grammar agent will find something here.",
        "Finally, the third paragraph concludes the document. It's short and simple."
    ]
    
    for p_text in paragraphs:
        doc.add_paragraph(p_text)
        
    doc.save(filename)
    print(f"Created {filename} at {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_dummy_docx()
